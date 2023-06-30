import sys
import datetime
from PyQt5.QtWidgets import QApplication, QMainWindow, QMessageBox, QPushButton, QLabel, QLineEdit, QVBoxLayout, QWidget, QFileDialog, QComboBox
from openpyxl import load_workbook
from docx import Document

class ProtocolGeneratorApp(QMainWindow):
    def __init__(self):
        super(ProtocolGeneratorApp, self).__init__()
        self.setWindowTitle('Protocol Generator')

        # Создание виджетов для интерфейса
        self.txtFilePath = QLineEdit(self)  # Поле для отображения выбранного пути к файлу
        self.txtFilePath.setReadOnly(True)
        self.btnLoadFile = QPushButton('Загрузить файл', self)  # Кнопка для загрузки файла
        self.comboOrganizations = QComboBox(self)  # Выпадающий список для выбора организации
        self.protocolNumber = QLineEdit(self)  # Поле для ввода номера протокола
        self.ProgramName = QLineEdit(self)  # Поле для ввода названия программы обучения
        self.WorkOuhrs = QLineEdit(self)  # Поле для ввода кол-ва рабочих часов
        self.txtChairperson = QLineEdit(self)  # Поле для ввода ФИО председателя комиссии
        self.txtNumberOfMembers = QLineEdit(self)  # Поле для ввода кол-ва членов комиссии
        self.btnConfirmMembers = QPushButton('Ок', self)  # Кнопка для подтверждения кол-ва членов комиссии
        self.btnGenerate = QPushButton('Сгенерировать протокол', self)  # Кнопка для генерации протокола

        # Создание центрального виджета и компоновка элементов интерфейса
        central_widget = QWidget(self)
        layout = QVBoxLayout(central_widget)
        layout.addWidget(QLabel('Путь к файлу:'))
        layout.addWidget(self.txtFilePath)
        layout.addWidget(self.btnLoadFile)
        layout.addWidget(QLabel('Организация:'))
        layout.addWidget(self.comboOrganizations)
        layout.addWidget(QLabel('Введите № протокола:'))
        layout.addWidget(self.protocolNumber)
        layout.addWidget(QLabel('Введите название программы обучения:'))
        layout.addWidget(self.ProgramName)
        layout.addWidget(QLabel('Введите кол-во рабочих часов:'))
        layout.addWidget(self.WorkOuhrs)
        layout.addWidget(QLabel('Введите кол-во членов комиссии'))
        layout.addWidget(self.txtNumberOfMembers)
        layout.addWidget(self.btnConfirmMembers)
        layout.addWidget(QLabel('ФИО председателя комиссии:'))
        layout.addWidget(self.txtChairperson)
        layout.addWidget(self.btnGenerate)

        self.setCentralWidget(central_widget)

        # Подключение обработчиков событий
        self.btnLoadFile.clicked.connect(self.load_file)
        self.btnGenerate.clicked.connect(self.generate_protocol)
        self.btnConfirmMembers.clicked.connect(self.load_members)

    def load_members(self):
        # Создание полей ввода на основе опроса пользователя
        layout = self.centralWidget().layout()  # Получаем макет виджета

        generated_fields = []  # Список для хранения сгенерированных полей

        for i in range(int(self.txtNumberOfMembers.text())):
            label = QLabel(f"ФИО члена комиссии {i + 1}:")
            line_edit = QLineEdit()

            # Вставляем поле для ввода перед кнопкой "Сгенерировать протокол"
            layout.insertWidget(layout.count() - 1, label)
            layout.insertWidget(layout.count() - 1, line_edit)

            generated_fields.append((label, line_edit))  # Сохраняем ссылки на сгенерированные поля

        # Сохраняем ссылки на сгенерированные поля ввода в объекте класса
        self.generated_fields = generated_fields

    def load_file(self):
        # Открытие диалогового окна для выбора файла
        file_dialog = QFileDialog(self)
        file_dialog.setWindowTitle('Выберите файл')
        file_dialog.setNameFilter('Файлы Excel (*.xlsx)')
        if file_dialog.exec_() == QFileDialog.Accepted:
            file_path = file_dialog.selectedFiles()[0]
            self.txtFilePath.setText(file_path)
            self.load_organizations(file_path)

    def load_organizations(self, file_path):
        # Загрузка организаций из файла и добавление их в выпадающий список
        self.comboOrganizations.clear()
        try:
            workbook = load_workbook(file_path)
            sheet = workbook.active
            organizations = set()
            for row in sheet.iter_rows(min_row=2, values_only=True):
                institution = row[2]  # 3-я колонка: Организация
                if institution is not None:  # Проверка на None
                    organizations.add(institution)
            self.comboOrganizations.addItems(sorted(organizations))

        except Exception as e:
            QMessageBox.critical(self, 'Ошибка', f'Произошла ошибка: {str(e)}', QMessageBox.Ok)

    def generate_protocol(self):
        file_path = self.txtFilePath.text()
        if not file_path:
            QMessageBox.critical(self, 'Ошибка', 'Выберите файл', QMessageBox.Ok)
            return

        organization = self.comboOrganizations.currentText()
        if not organization:
            QMessageBox.critical(self, 'Ошибка', 'Выберите организацию', QMessageBox.Ok)
            return

        # Следующие проверки по идее можно убрать, чтобы поля пустыми могли оставаться
        chairperson = self.txtChairperson.text()
        if not chairperson:
            QMessageBox.critical(self, 'Ошибка', 'Введите ФИО председателя комиссии', QMessageBox.Ok)
            return

        members = []
        for label, line_edit in self.generated_fields:
            member_name = line_edit.text()
            if not member_name:
                QMessageBox.critical(self, 'Ошибка', f'Введите ФИО члена комиссии {len(members) + 1}', QMessageBox.Ok)
                return
            members.append(member_name)

        try:
            workbook = load_workbook(file_path)
            sheet = workbook.active

            document = Document('shablon_praktika.docx') # Путь до шаблона
            names = []
            positions = []
            institution2 = None
            protocol_number = None

            # Поиск и сбор данных для выбранной организации
            for row in sheet.iter_rows(min_row=2, values_only=True):
                institution = row[2]  # 3-я колонка: Организация

                if institution is not None and institution == organization:  # Проверка на None и сравнение
                    name = row[3]  # 4-я колонка: ФИО
                    position = row[4]  # 5-я колонка: Должность
                    institution2 = row[9]  # 10-я колонка: Организация-2
                    protocol_number = row[10]  # 11-я колонка: Номер протокола
                    names.append(name)
                    positions.append(position)

            # Генерация протокола
            for paragraph in document.paragraphs:
                if '{{Program}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{{Program}}', self.ProgramName.text())
                if '{{hours}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{{hours}}', self.WorkOuhrs.text())
                if '{{organization}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{{organization}}', organization)
                if '{{chairperson}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{{chairperson}}', chairperson)

                # Замена заполнителя {{members}}
                if '{{members}}' in paragraph.text:
                    members_list = ', '.join(members)
                    paragraph.text = paragraph.text.replace('{{members}}', members_list)

                if '{{institution2}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{{institution2}}', institution2)
                if '{{protocol_number}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{{protocol_number}}', self.protocolNumber.text())


            # Вставка данных из таблицы
            table = document.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            header_cells[0].text = '№'
            header_cells[1].text = 'ФИО'
            header_cells[2].text = 'Должность'
            header_cells[3].text = 'Результат проверки знаний'
            header_cells[4].text = 'Регистрационный номер записи'
            header_cells[5].text = 'Подпись проверяемого'

            cnt = 0
            for name, position in zip(names, positions):
                cnt += 1
                row_cells = table.add_row().cells
                row_cells[0].text = str(cnt) + "."
                row_cells[1].text = name
                row_cells[2].text = position

            document.add_paragraph("")
            # Вставка строк "Члены комиссии"
            document.add_paragraph("Председатель комиссии:\t____________ " + chairperson)
            for member in members:
                document.add_paragraph("Члены комиссии:\t\t____________ " + member)
                break
            if len(members) > 1:
                for i in range(1, len(members)):
                    document.add_paragraph("\t\t\t\t____________ " + members[i])

            # Определение текущей даты
            current_date = datetime.date.today()
            formatted_date = current_date.strftime('%d %m %Y') + ' года'

            # Замена заполнителя {{current_date}}
            for paragraph in document.paragraphs:
                if '{{current_date}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{{current_date}}', formatted_date)
            file_dialog = QFileDialog(self)
            file_dialog.setWindowTitle('Сохранить файл')
            file_dialog.setAcceptMode(QFileDialog.AcceptSave)
            file_dialog.setNameFilter('Документы Word (*.docx)')

            # Предустановленное название файла
            default_file_name = organization
            # Начальное имя файла
            file_dialog.selectFile(default_file_name)

            if file_dialog.exec_() == QFileDialog.Accepted:
                file_path = file_dialog.selectedFiles()[0]
                document.save(file_path) # Путь сохранения выходного файла
            QMessageBox.information(self, 'Успех', 'Протокол успешно сгенерирован', QMessageBox.Ok)

        except Exception as e:
            QMessageBox.critical(self, 'Ошибка', f'Произошла ошибка: {str(e)}', QMessageBox.Ok)

if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = ProtocolGeneratorApp()
    window.show()
    sys.exit(app.exec_())